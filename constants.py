import ruamel.yaml

'''
Constants - Reads text file and loads the constants in them
Including allowed ranges for tests, document template names, etc.
'''
class constants():
   '''
   Initializes instance variables and sets them by reading the yaml
   Called when object is created
   '''   
   def __init__(self) -> None: 
      # initializing constants 
      self.PARAMETERS:list
      self.dictSymbol:dict
      # template names
      self.COA_TEMPLATE_NAME:str
      self.LABSHEET_TEMPLATE_NAME:str
      self.COA_FAILED_TEMPLATE_NAME:str
      self.LABSHEET_FAILED_TEMPLATE_NAME:str
      # allowed ranges
      self.allowedRangesI:list
      self.allowedRangesII:list
      self.allowedRangesIII:list

      self.loadYaml()


   '''
   Load all constants from config.yaml
   '''
   def loadYaml(self) -> None:
      # open and read config.yaml
      with open("config.yaml", "r", encoding="utf-8") as file:
         # stream:dict # type hinting dictionary
         stream = ruamel.yaml.safe_load(file) # safely load the yaml file

      # setting constants   
      self.PARAMETERS = stream.get('PARAMETERS')
      self.dictSymbol = stream.get('dictSymbol')
      # template names
      self.COA_TEMPLATE_NAME = stream.get('COA_TEMPLATE_NAME')
      self.LABSHEET_TEMPLATE_NAME = stream.get('LABSHEET_TEMPLATE_NAME')
      self.COA_FAILED_TEMPLATE_NAME = stream.get('COA_FAILED_TEMPLATE_NAME')
      self.LABSHEET_FAILED_TEMPLATE_NAME = stream.get('LABSHEET_FAILED_TEMPLATE_NAME')
      # allowed ranges
      self.allowedRangesI = stream.get('allowedRangesI')
      self.allowedRangesII = stream.get('allowedRangesII')
      self.allowedRangesIII = stream.get('allowedRangesIII')


   def updateTemplateRange(self) -> None:
      try:
         import docx

         COA_doc = docx.Document('templates/' + self.COA_TEMPLATE_NAME)
         COA_F_doc = docx.Document('templates/' + self.COA_FAILED_TEMPLATE_NAME)
         Lab_doc = docx.Document('templates/' + self.LABSHEET_TEMPLATE_NAME)
         Lab_F_doc = docx.Document('templates/' + self.LABSHEET_FAILED_TEMPLATE_NAME)
         for r in range(1, 4):
            for c in range(1, 11):
               writeCell(COA_doc, r, c, self.rangeToText(r, c, True), True)
               writeCell(COA_F_doc, r, c, self.rangeToText(r, c, True), True)
               writeCell(Lab_doc, r, c, self.rangeToText(r, c, False), False)
               writeCell(Lab_F_doc, r, c, self.rangeToText(r, c, False), False)

         COA_doc.save('templates/' + self.COA_TEMPLATE_NAME)
         COA_F_doc.save('templates/' + self.COA_FAILED_TEMPLATE_NAME)
         Lab_doc.save('templates/' + self.LABSHEET_TEMPLATE_NAME)
         Lab_F_doc.save('templates/' + self.LABSHEET_FAILED_TEMPLATE_NAME)

      except Exception as e:
         if type(e) == PermissionError: raise Exception("Please close all opened templates")
         else: raise e


   '''
   (Row and col of the table)
   '''
   def rangeToText(self, row:int, col:int, isCOA:bool) -> str:
      if row == 1:
         ranges = self.allowedRangesI
      elif row == 2:
         ranges = self.allowedRangesII
      elif row == 3:
         ranges = self.allowedRangesIII
      else:
         raise IndexError("Row index out of range")

      # get the range (one range)
      r = ranges[col-1].copy()

      if isCOA:
         # add additional edits for COA
         for i in range(len(r)):
            for s, replacement in {"TRA":"Trace", "(-)":"NEG (-)",  "(+)":"POS (+)"}.items():
               r[i] = r[i].replace(s, replacement)

         # adding units where necessary
         if col == 3: # URO
            r[-1] = r[-1] + " EU/dL" 
         elif col in [4, 8, 10]: # PRO, KET, GLU
            if r[-1] != "NEG (-)":
               r[-1] = r[-1] + " mg/dL" 

         # add 'to' when appropriate
         if len(r) > 1:
            return r[0] + " to " + r[-1]
         else:
            return r[0]
      else:
         # add '-' when appropriate
         if len(r) > 1:
            return r[0] + " - " + r[-1]
         else:
            return r[0]


   '''
   Updates constants by reading the template for ranges,
   write to the config file, then load it to the program
   '''
   def updateConstants(self) -> None:
      self.updateYaml(self.readTemplateRanges())
      self.loadYaml()


   '''
   Read COA Template docx and write new ranges onto yaml, then reload the yaml

   Args:
         None
   
   Returns:
         list: List of 30 lists of allowed ranges (in str) read from COA Template
   '''
   def readTemplateRanges(self) -> list:
      from docx2python import docx2python
      # extract docx content
      doc_result = docx2python('templates/' + self.COA_TEMPLATE_NAME).body
      # extract tolerence range and store them in one long list
      doc_result = doc_result[7]
      templateRanges = doc_result[1][1:11]
      templateRanges.extend(doc_result[2][1:11])
      templateRanges.extend(doc_result[3][1:11])

      # DEBUG: print(templateRanges) # is list[list[str]] 

      processedRanges = []
      for range in templateRanges:
         # take it out of the weird list format
         range = range[0].split(' ')

         # DEBUG: range = ['sample ranges here']
         # if range is actually a range
         if 'to' in range:
            # for BIL, when NEG is in range too
            if range[0] == 'NEG':
               # adds (-) then generate the list of ranges from 1+ to end
               processedRanges.append(["(-)"] + generateInterval('1+', range[3]))
            else:
               # else, generateInterval()
               processedRanges.append(generateInterval(range[0], range[2]))
         else: 
            # when not range, it's a +/-, so tae range[1] (where the symbol is)
            processedRanges.append([range[1]])
      
      # DEBUG: print(processedRanges)
      return processedRanges

      
   '''
   Updates config.yaml based on new constants

   Args:
         newConfig (list): List of 30 lists of allowed ranges (in str), from readTemplateRanges()

   Returns:
         None
   '''
   def updateYaml(self, newConfig:list) -> None:
      yaml = ruamel.yaml.YAML()

      # opens yaml file while preserving encoding & quotes
      with open('config.yaml', encoding="utf-8") as fp:
         yaml.preserve_quotes = True
         data = yaml.load(fp)
      
      # turning all lists in newConfig into ruamel.yaml.CommentSeq
      # allows us to set flow style to output yaml
      l = []
      for i in newConfig:
         l.append(seq(i))

      # update ranges based on appropriate pos of ranges
      data.update({'allowedRangesI': l[0:10]})
      data.update({'allowedRangesII': l[10:20]})
      data.update({'allowedRangesIII': l[20:30]})

      # more yaml formatting
      yaml.indent(mapping=2, sequence=3, offset=1)
      
      # write the updated data back onto the yaml file
      with open('config.yaml', 'w', encoding="utf-8") as fp:
         yaml.dump(data, fp)


'''
Generates interval of allowed ranges from start to end
Not very smart tho

Args:
      start (str): starting range
      end (str): sending range

Returns:
      list[str]: the list of all possible ranges given the start & end
'''
def generateInterval(start:str, end:str) -> list[str]:
   try:
      # for 1+/2+/3+ format (LEU, BLO, BIL)
      if start in ['1+', '2+', '3+']:
         outputRange = []
         for i in range(int(start[0]), int(end[0]) + 1):
            outputRange.append(str(i) + '+')
         return outputRange

      # for URO
      if start in ['0.2', '1.0', '4.0']:
         return [start, end]

      # for PRO
      if start in ['100', 'Trace']:
         return ['TRA' if s=='Trace' else s for s in [start, end]]

      # for pH
      if start in ['5.0', '5.5', '6.0', '6.5', '7.0', '7.5', '8.0', '8.5']:
         start = int(float(start) * 2)
         end = int(float(end) * 2)
         return [str(i/2) for i in range(start, end + 1)]
      
      # for SG
      if start in ['1.000', '1.005', '1.010', '1.015', '1.020']:
         start = int(start[3:5])
         end = int(end[3:5])
         return ['1.0' + str(i) for i in range(start, end + 5, 5)]
      
      # for KET - TODO: add 15
      if start in ['40', '\u226580']:
         return [start, end]

      # for GLU
      if start in ['100', '500']:
         return [start, end]

      raise Exception()
   except:
      print("Ranges from template not included in code")
      raise Exception("Ranges from template not included in code")


'''
Helper func for updateYaml,
Neccessary for setting output yaml in flow style (instead of block)

Args:
      l (list[str]): 

Returns:
      CommentedSeq: The ramual.yaml compatible version of the list
'''
def seq(l:list[str]) -> ruamel.yaml.CommentedSeq:
   s = ruamel.yaml.comments.CommentedSeq(l)
   s.fa.set_flow_style()
   return s


'''
Write text to designated cell and set appropriate formatting
'''
def writeCell(doc, row:int, col:int, text:str, isCOA:bool) -> None:
   from docx.enum.text import WD_ALIGN_PARAGRAPH

   # get the table
   if isCOA:      
      table = doc.tables[3]
   else:
      table = doc.tables[6]
   # get cell 
   cell = table.cell(row, col)
   # set cell text
   cell.text = text
   # get paragraph
   paragraph = cell.paragraphs[0]
   # set text alignment
   paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
   # set text font
   run = paragraph.runs
   run[0].font.name = 'Arial'   


# for debug use
if __name__ == '__main__':
   c = constants()
   # c.updateConstants()
   c.updateTemplateRange()